import re
import os
import spacy
from spacy.matcher import PhraseMatcher
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from pypdf import PdfReader
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# List of tech skills compiled across frontend, backend, mobile, data science, ML, devops, database, languages
SKILL_TAXONOMY = [
    # Programming Languages
    "javascript", "typescript", "python", "java", "c++", "go", "rust", "ruby", "php", "swift", "kotlin", "sql", "html", "css", "r", "scala", "shell", "bash",
    # Frontend
    "react", "angular", "vue", "next.js", "nextjs", "nuxt", "gatsby", "redux", "svelte", "html5", "css3", "sass", "less", "tailwind", "bootstrap", "material ui", "webpack", "vite",
    # Backend & Frameworks
    "node.js", "nodejs", "express", "express.js", "django", "flask", "fastapi", "spring boot", "spring", "hibernate", "laravel", "ruby on rails", "rails", "asp.net", "dotnet", "net core", "nestjs", "graphql", "rest api", "restful api", "microservices", "websockets", "grpc",
    # Mobile
    "react native", "flutter", "xamarin", "objective-c", "android studio", "ios", "android",
    # Databases
    "mongodb", "postgresql", "postgres", "mysql", "redis", "cassandra", "sqlite", "dynamodb", "oracle", "sql server", "firebase", "mariadb", "neo4j",
    # DevOps, Cloud & CI/CD
    "docker", "kubernetes", "k8s", "jenkins", "git", "github", "github actions", "gitlab", "aws", "amazon web services", "azure", "gcp", "google cloud", "terraform", "ansible", "linux", "unix", "nginx", "apache", "ci/cd", "cicd", "prometheus", "grafana", "elk stack", "nagios",
    # ML, Data Science & AI
    "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "pytorch", "keras", "opencv", "nltk", "spacy", "huggingface", "transformers", "deep learning", "machine learning", "artificial intelligence", "data visualization", "tableau", "powerbi", "power bi", "matplotlib", "seaborn", "statistics", "probability", "big data", "hadoop", "spark", "pyspark", "nlp", "natural language processing", "computer vision", "llm", "langchain", "rag",
    # Other Tools & Methodologies
    "agile", "scrum", "jira", "confluence", "system design", "data structures", "algorithms", "dsa", "unit testing", "jest", "mocha", "cypress", "postman"
]

# Supported job roles definitions
ROLE_TEMPLATES = {
    "Full Stack Developer": {
        "skills": ["react", "node.js", "nodejs", "javascript", "express", "express.js", "html", "css", "mongodb", "sql", "git", "rest api", "typescript"],
        "roadmap": [
            {"step": 1, "title": "TypeScript Mastery", "description": "Learn advanced TypeScript types, interfaces, generics, and compiler options.", "resources": "TypeScript Official HandBook, Execute TypeScript with Node"},
            {"step": 2, "title": "Next.js & SSR", "description": "Master Next.js App Router, Server Actions, Server-Side Rendering (SSR), and Static Site Generation (SSG).", "resources": "Next.js official tutorial, Vercel documentation"},
            {"step": 3, "title": "Database Optimization", "description": "Learn SQL indexing, MongoDB aggregation framework, transactions, and caching with Redis.", "resources": "MongoDB University, High-Performance MySQL"},
            {"step": 4, "title": "System Design & Scalability", "description": "Study vertical vs horizontal scaling, load balancing, message queues (RabbitMQ/Kafka), and microservices.", "resources": "System Design Primer, ByteByteGo"}
        ],
        "recommendedProjects": [
            {"title": "Real-time Collaborative Board", "description": "A collaborative canvas tool utilizing WebSockets, React, Node.js, and Redis for instant updates.", "techStack": "React, Node.js, Socket.io, Redis"},
            {"title": "Microservices E-commerce", "description": "An e-commerce platform split into Auth, Product, and Order services with Docker, RabbitMQ, and MongoDB.", "techStack": "Express, MongoDB, RabbitMQ, Docker, Gateway"}
        ],
        "recommendedCertifications": [
            "AWS Certified Developer - Associate",
            "Meta Back-End Developer Professional Certificate",
            "MongoDB Certified Developer Associate"
        ]
    },
    "Data Scientist": {
        "skills": ["python", "sql", "pandas", "numpy", "scikit-learn", "sklearn", "statistics", "probability", "data visualization", "tableau", "powerbi", "power bi", "matplotlib", "seaborn", "spark"],
        "roadmap": [
            {"step": 1, "title": "Advanced Statistics & Hypothesis Testing", "description": "Deepen your knowledge on A/B testing, regression analysis, ANOVA, and probability distributions.", "resources": "Khan Academy Statistics, OpenIntro Statistics"},
            {"step": 2, "title": "SQL & ETL Pipelines", "description": "Master complex subqueries, window functions, and setting up ETL pipelines using Airflow.", "resources": "SQL Zoo, Mode Analytics SQL Tutorial"},
            {"step": 3, "title": "Machine Learning Engineering", "description": "Learn tuning hyper-parameters, cross-validation, feature engineering, and model evaluation metrics.", "resources": "Scikit-Learn documentation, Hands-On Machine Learning book"},
            {"step": 4, "title": "Big Data Technologies", "description": "Understand distributed computing basics using Apache Spark/PySpark.", "resources": "Databricks Academy Spark tutorial"}
        ],
        "recommendedProjects": [
            {"title": "Customer Churn Prediction Engine", "description": "Analyze subscription data, build classification models (XGBoost, Random Forest), evaluate using ROC-AUC, and write a report.", "techStack": "Python, Pandas, Scikit-learn, Seaborn"},
            {"title": "ETL Sales Dashboard", "description": "Build an automated ETL pipeline that loads raw CSVs into PostgreSQL and visualizes insights on an interactive dashboard.", "techStack": "SQL, PostgreSQL, Python, Airflow, Tableau"}
        ],
        "recommendedCertifications": [
            "IBM Data Science Professional Certificate",
            "Google Data Analytics Professional Certificate",
            "Microsoft Certified: Power BI Data Analyst Associate"
        ]
    },
    "ML Engineer": {
        "skills": ["python", "tensorflow", "pytorch", "machine learning", "deep learning", "git", "docker", "mlops", "sql", "scikit-learn", "sklearn", "pandas", "numpy", "nlp", "computer vision"],
        "roadmap": [
            {"step": 1, "title": "Deep Learning Foundation", "description": "Learn neural network theory, backpropagation, CNNs, RNNs, and Transformers.", "resources": "Deep Learning Specialization by Andrew Ng, PyTorch Tutorials"},
            {"step": 2, "title": "MLOps & Model Deployment", "description": "Build APIs using FastAPI to serve models, run inference in Docker containers, and set up CI/CD.", "resources": "MLOps Guide, Made With ML"},
            {"step": 3, "title": "Model Optimization & Quantization", "description": "Learn models pruning, quantization, ONNX runtime conversion to speed up inference.", "resources": "PyTorch Quantization guide"},
            {"step": 4, "title": "Kubernetes & Scalable ML", "description": "Orchestrate multi-container ML pipelines with Kubeflow on Kubernetes.", "resources": "Kubeflow documentation"}
        ],
        "recommendedProjects": [
            {"title": "Real-time Object Detection API", "description": "Deploy a YOLO model using PyTorch and FastAPI, with optimized Docker configuration for fast inference.", "techStack": "Python, PyTorch, OpenCV, FastAPI, Docker"},
            {"title": "Fine-Tuning LLM for Q&A", "description": "Fine-tune a small LLM (like Llama-3 or Mistral) on a custom dataset using PEFT/LoRA and Hugging Face.", "techStack": "Python, HuggingFace, PyTorch, Transformers, PEFT"}
        ],
        "recommendedCertifications": [
            "Google Cloud Professional Machine Learning Engineer",
            "AWS Certified Machine Learning - Specialty",
            "TensorFlow Developer Certificate"
        ]
    },
    "Java Developer": {
        "skills": ["java", "spring boot", "spring", "hibernate", "sql", "rest api", "git", "maven", "microservices", "junit", "postgresql", "mysql"],
        "roadmap": [
            {"step": 1, "title": "Java Concurrency & Internals", "description": "Learn JVM tuning, garbage collection mechanisms, multithreading, and Java concurrency API.", "resources": "Java Concurrency in Practice, Baeldung Java"},
            {"step": 2, "title": "Spring Boot Microservices", "description": "Build resilient architectures with Spring Cloud, Eureka Discovery, API Gateway, and Feign Clients.", "resources": "Spring Cloud Guides, Java Brains Microservices"},
            {"step": 3, "title": "Security & OAuth2", "description": "Secure Spring Boot REST APIs using JWT tokens, Spring Security, and OAuth2 integration.", "resources": "Spring Security documentation"},
            {"step": 4, "title": "Containerization & Deployment", "description": "Dockerize Java applications, create multi-stage builds, and deploy on AWS or Kubernetes.", "resources": "Docker official guide for Spring Boot"}
        ],
        "recommendedProjects": [
            {"title": "Distributed Banking Application", "description": "A microservices banking system using Spring Cloud Config, Eureka registry, API Gateway, and PostgreSQL.", "techStack": "Java, Spring Boot, Spring Cloud, PostgreSQL, Docker"},
            {"title": "High-Throughput Order Management", "description": "An order processor using Kafka queues, Hibernate ORM caching, Redis cache, and JUnit/Mockito tests.", "techStack": "Java, Spring Boot, Kafka, Redis, Hibernate"}
        ],
        "recommendedCertifications": [
            "Oracle Certified Professional: Java SE Developer",
            "Spring Certified Professional"
        ]
    },
    "DevOps Engineer": {
        "skills": ["docker", "kubernetes", "linux", "unix", "jenkins", "aws", "gcp", "terraform", "git", "github actions", "ci/cd", "cicd", "ansible", "prometheus", "grafana", "nginx", "bash", "python"],
        "roadmap": [
            {"step": 1, "title": "Linux Administration & Bash Scripting", "description": "Master command line, process management, networking troubleshooting, and automation scripting.", "resources": "Linux Journey, Advanced Bash Scripting Guide"},
            {"step": 2, "title": "Infrastructure as Code (IaC)", "description": "Learn Terraform for cloud resource management on AWS, workspace modularization, and state management.", "resources": "Terraform Up & Running, HashiCorp Learn"},
            {"step": 3, "title": "Kubernetes Administration", "description": "Understand Kubernetes architecture, Pods, Services, Deployments, ConfigMaps, Secrets, Ingress, and Helm.", "resources": "Mumshad's CKA course, Kubernetes docs"},
            {"step": 4, "title": "Monitoring, Alerting & Logging", "description": "Set up Prometheus for metric gathering, Grafana for dashboard visualizations, and Loki/ELK for logging.", "resources": "Prometheus & Grafana official docs"}
        ],
        "recommendedProjects": [
            {"title": "Auto-scaling AWS Infrastructure", "description": "Deploy a secure VPC with private and public subnets, ASG, Load Balancers, and RDS using Terraform.", "techStack": "Terraform, AWS, VPC, EC2, RDS"},
            {"title": "Multi-stage CI/CD Pipeline", "description": "Build a pipeline in GitHub Actions/Jenkins that tests a Node app, builds a Docker image, pushes to ECR, and deploys to EKS.", "techStack": "Jenkins, GitHub Actions, Docker, AWS EKS, Slack Notification"}
        ],
        "recommendedCertifications": [
            "AWS Certified DevOps Engineer - Professional",
            "Certified Kubernetes Administrator (CKA)",
            "HashiCorp Certified: Terraform Associate"
        ]
    }
}

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract all text page-by-page from a PDF file."""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a Word (.docx) file paragraphs and tables."""
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
        return ""

def extract_text(file_path: str) -> str:
    """Unified function to extract text based on file extension."""
    if not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        # Try reading as text file
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

def clean_text(text: str) -> str:
    """Preprocess and clean the text corpus for similarity checking."""
    # Remove URLs
    text = re.sub(r'http\S+\s*', ' ', text)
    # Remove RT and cc
    text = re.sub(r'RT|cc', ' ', text)
    # Remove hashtags
    text = re.sub(r'#\S+', ' ', text)
    # Remove email references
    text = re.sub(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', ' ', text)
    # Remove phone number references
    text = re.sub(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', ' ', text)
    # Remove special characters except common punctuation/spaces
    text = re.sub(r'[^\w\s\.\,\-\#\+\/\(\)]', ' ', text)
    # Replace multiple spaces with single space
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()

def extract_contact_info(text: str):
    """Extract email, phone, and name from the resume text using regex and spaCy."""
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'((?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+?\d{10,12})'
    
    email_match = re.search(email_pattern, text)
    phone_match = re.search(phone_pattern, text)
    
    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0) if phone_match else ""
    
    # Post-process email to strip common prefix labels (e.g. "Envelope" or "email")
    if email:
        for prefix in ["envelope", "email", "mailto", "contact", "address", "phone"]:
            if email.lower().startswith(prefix):
                email = email[len(prefix):]
                
    # Blacklisted words that are definitely not PERSON names
    BLACKLIST = {
        "linkedin", "github", "resume", "curriculum", "vitae", "cv", 
        "page", "email", "phone", "mobile", "address", "summary", 
        "experience", "education", "skills", "projects", "certifications",
        "developer", "engineer", "designer", "manager", "portfolio"
    }
    
    def is_valid_name(name_str: str) -> bool:
        name_lower = name_str.lower()
        if not re.match(r'^[a-zA-Z\s\-]+$', name_str):
            return False
        # If any word in the name is in blacklist, it's invalid
        words = name_lower.split()
        for w in words:
            if w in BLACKLIST:
                return False
        return True

    # Simple Name extraction heuristic
    # Look at spaCy NER PERSON tag first in the top portion of the resume
    doc = nlp(text[:800])
    name = ""
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            val = ent.text.strip()
            # Clean up newlines or weird chars in name
            val = re.sub(r'\s+', ' ', val)
            # Make sure it contains letters, reasonable length, and is not blacklisted
            if len(val) > 2 and is_valid_name(val):
                name = val
                break
    
    # Fallback name extraction: look for labeled "Name:" lines or title-case first lines
    if not name:
        lines = [line.strip() for line in text.split("\n") if line.strip()]

        # 1) Look for explicit labels like 'Name: John Doe'
        for line in lines[:10]:
            m = re.search(r'^(?:name)\s*[:\-\s]{1,3}(.+)$', line, flags=re.I)
            if m:
                candidate = m.group(1).strip()
                candidate = re.sub(r'\s+', ' ', candidate)
                if 1 < len(candidate.split()) < 5 and is_valid_name(candidate):
                    name = candidate
                    break

        # 2) If still not found, try a Title Case heuristic on the top lines (2-3 words, capitalized)
        if not name:
            for line in lines[:10]:
                # ignore lines that are clearly headers or contain '@' or digits
                if '@' in line or any(char.isdigit() for char in line):
                    continue
                # Candidate should be short (2-4 words) and mostly Title Case words
                words = line.split()
                if 1 < len(words) <= 4:
                    title_like = sum(1 for w in words if re.match(r'^[A-Z][a-z\-]+$', w))
                    if title_like >= max(1, len(words) - 1) and is_valid_name(line):
                        name = line
                        break

    # If not found, leave name empty so frontend shows 'N/A' instead of a placeholder
    if not name:
        name = ""

    # Also try to extract email/phone from labeled lines if initial regex missed them
    if not email:
        for line in text.split("\n")[:20]:
            m = re.search(r'(?:email|e-mail|e:)\s*[:\-\s]{0,3}([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', line, flags=re.I)
            if m:
                email = m.group(1).strip()
                break

    if not phone:
        for line in text.split("\n")[:20]:
            m = re.search(r'(?:phone|mobile|tel|contact)\s*[:\-\s]{0,3}((?:\+?\d[\d\s\-\(\)]{6,}\d))', line, flags=re.I)
            if m:
                phone = m.group(1).strip()
                break

    return {
        "name": name,
        "email": email,
        "phone": phone
    }


def extract_skills(text: str) -> list:
    """Extract skills from the resume text matching against our SKILL_TAXONOMY."""
    cleaned = clean_text(text)
    # Tokenize the text to avoid partial word matching (e.g. matching "c" in "docker")
    # For compound skills (e.g. "spring boot", "react native"), we use direct string check or spaCy phrase matching
    found_skills = []
    
    # We do a fast phrase matching using lowercase
    # To prevent boundary issues (like matching "go" inside "good"), we use regex boundaries for word characters
    for skill in SKILL_TAXONOMY:
        # Escape any special characters in the skill (like c++, .net, next.js)
        escaped_skill = re.escape(skill)
        # Regex check with word boundaries (handling special chars at boundaries too)
        # If it starts or ends with symbols like +, ., we have to handle boundaries differently
        pattern = rf'\b{escaped_skill}\b'
        if skill.endswith('+') or skill.startswith('.') or skill.endswith('.'):
            # Relax boundary check for special characters
            pattern = rf'(?:^|\s){escaped_skill}(?:$|\s|[.,;])'
            
        if re.search(pattern, cleaned):
            # Normalize skills names
            normalized = skill
            if skill == "nodejs": normalized = "node.js"
            if skill == "nextjs": normalized = "next.js"
            if skill == "expressjs": normalized = "express.js"
            if skill == "sklearn": normalized = "scikit-learn"
            if skill == "postgres": normalized = "postgresql"
            if skill == "k8s": normalized = "kubernetes"
            if skill == "cicd": normalized = "ci/cd"
            
            if normalized not in found_skills:
                found_skills.append(normalized)
                
    return found_skills

def extract_education_experience(text: str):
    """Extract education, experience, certifications, and projects using keyword section parsing."""
    lines = text.split("\n")
    sections = {
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": []
    }
    
    current_section = None
    
    edu_keywords = ["education", "academic", "degree", "university", "college", "school"]
    exp_keywords = ["experience", "work history", "employment", "professional experience", "career history", "history"]
    proj_keywords = ["projects", "personal projects", "academic projects", "key projects"]
    cert_keywords = ["certifications", "licenses", "certificates", "courses", "credentials"]
    
    for line in lines:
        cleaned_line = line.strip().lower()
        if not cleaned_line:
            continue
            
        # Check if line is a section header (usually short, <= 4 words, containing keywords)
        words = cleaned_line.split()
        if len(words) <= 4:
            matched = False
            for kw in edu_keywords:
                if kw in cleaned_line:
                    current_section = "education"
                    matched = True
                    break
            if not matched:
                for kw in exp_keywords:
                    if kw in cleaned_line:
                        current_section = "experience"
                        matched = True
                        break
            if not matched:
                for kw in proj_keywords:
                    if kw in cleaned_line:
                        current_section = "projects"
                        matched = True
                        break
            if not matched:
                for kw in cert_keywords:
                    if kw in cleaned_line:
                        current_section = "certifications"
                        matched = True
                        break
            if matched:
                continue
                
        # Append line to corresponding section
        if current_section and len(line.strip()) > 3:
            sections[current_section].append(line.strip())
            
    # Format and truncate long sections for DB storage
    return {
        "education": "\n".join(sections["education"][:15]) if sections["education"] else "Not explicitly detailed.",
        "experience": "\n".join(sections["experience"][:25]) if sections["experience"] else "Not explicitly detailed.",
        "projects": "\n".join(sections["projects"][:20]) if sections["projects"] else "Not explicitly detailed.",
        "certifications": "\n".join(sections["certifications"][:15]) if sections["certifications"] else "Not explicitly detailed."
    }

def analyze_resume_ats(resume_text: str, target_role: str, resume_skills: list):
    """Calculate ATS compatibility score and gap analysis."""
    # Fallback to Full Stack Developer template if role is not recognized
    role = target_role if target_role in ROLE_TEMPLATES else "Full Stack Developer"
    template = ROLE_TEMPLATES[role]
    required_skills = template["skills"]
    
    # 1. Skill Relevance Score (based on intersection)
    matching_skills = [s for s in resume_skills if s in required_skills or any(rs in s or s in rs for rs in required_skills)]
    matching_skills = list(set(matching_skills)) # unique
    
    missing_skills = [s for s in required_skills if s not in resume_skills and not any(rs in s or s in rs for rs in resume_skills)]
    
    if len(required_skills) > 0:
        skill_relevance = round((len(matching_skills) / len(required_skills)) * 100)
    else:
        skill_relevance = 0
        
    # 2. Keyword Density & Context Similarity (using TF-IDF + Cosine Similarity)
    # Build a virtual job description document using target role required skills and roadmap
    jd_doc = " ".join(required_skills) + " " + " ".join([step["title"] + " " + step["description"] for step in template["roadmap"]])
    
    cleaned_resume = clean_text(resume_text)
    cleaned_jd = clean_text(jd_doc)
    
    try:
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([cleaned_resume, cleaned_jd])
        cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        keyword_density = round(cosine_sim * 100)
    except Exception:
        # Fallback to basic text overlap
        keyword_density = 50
        
    # Limit extreme bounds
    keyword_density = max(10, min(100, keyword_density))
    
    # 3. Structure Score
    # Give points for having standard sections and contact info
    structure_score = 0
    checks = []
    
    contact_info = extract_contact_info(resume_text)
    if contact_info["email"]:
        structure_score += 20
        checks.append("Email Contact Information present")
    else:
        checks.append("Missing Email details")
        
    if contact_info["phone"]:
        structure_score += 20
        checks.append("Phone Contact Information present")
    else:
        checks.append("Missing Phone details")
        
    cleaned_lower = resume_text.lower()
    if any(k in cleaned_lower for k in ["experience", "work history", "employment"]):
        structure_score += 20
        checks.append("Experience/Work History section parsed")
    else:
        checks.append("Missing clear Work History header")
        
    if any(k in cleaned_lower for k in ["education", "university", "college", "degree"]):
        structure_score += 20
        checks.append("Education/Academic history parsed")
    else:
        checks.append("Missing clear Education section")
        
    if any(k in cleaned_lower for k in ["projects", "certifications", "skills"]):
        structure_score += 20
        checks.append("Skills/Projects or Certifications section parsed")
    else:
        checks.append("Missing structured Skills or Projects section")
        
    # 4. Overall ATS Score
    # Weighted average: 45% Skill Relevance, 35% Keyword Similarity, 20% Structure
    overall_score = round((skill_relevance * 0.45) + (keyword_density * 0.35) + (structure_score * 0.20))
    overall_score = max(5, min(100, overall_score))
    
    # Improvement Suggestions
    suggestions = []
    if skill_relevance < 70:
        suggestions.append(f"Add missing core skills required for {role}: {', '.join(missing_skills[:4])}.")
    if keyword_density < 50:
        suggestions.append("Incorporate more keywords from the target job descriptions naturally into your project and experience descriptions.")
    if structure_score < 100:
        suggestions.append("Structure your resume sections with standard titles like 'Professional Experience', 'Education', and ensure email/phone are easily readable.")
    if len(matching_skills) < 5:
        suggestions.append("Detail more software engineering tools, databases, and programming languages you have hands-on experience with.")
        
    if not suggestions:
        suggestions.append("Your resume matches this role exceptionally well! Consider tailoring specific metrics and quantities in your job details.")
        
    # Interview Preparation Guide generator based on role and missing skills
    interview_prep = {
        "technical_questions": [
            f"Can you explain the core concepts of {s.upper()} and how you have used it?" for s in (matching_skills[:2] + required_skills[:2])[:3]
        ],
        "behavioral_guidance": "Prepare STAR-method answers (Situation, Task, Action, Result) highlighting problem solving, team collaboration, and dealing with technical debt. Explain your projects clearly, stressing on the 'Why' of the technology choices.",
        "role_specific_focus": f"For a {role} role, expect questions on System Architecture, debugging strategies, security best practices, and runtime optimization of your chosen stacks."
    }
    
    return {
        "score": {
            "overall": overall_score,
            "skillRelevance": skill_relevance,
            "keywordDensity": keyword_density,
            "structureScore": structure_score,
            "checks": checks
        },
        "matchingSkills": matching_skills,
        "missingSkills": missing_skills,
        "suggestions": suggestions,
        "roadmap": template["roadmap"],
        "recommendedProjects": template["recommendedProjects"],
        "recommendedCertifications": template["recommendedCertifications"],
        "interviewPrep": interview_prep
    }
